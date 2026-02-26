"""
Ingestion Pipeline - Master Function with Enhanced Latent Hierarchy Generation

This module provides the main `ingest_document()` function that implements
the full Latent Hierarchy Generator from the research paper (Section 4-6).

TIER 1: Visual-Geometric Analysis (Primary)
    1a. PyMuPDF Font Histogram (Section 6.1)
        - If headers detected via font variance → Build hierarchical tree
    1b. Recursive XY-Cut (Section 4.1.1) - Optional for complex layouts
        - For multi-column documents, L-shaped text wraps
        
TIER 2: Semantic Boundary Detection (Fallback 1 - Flat Text)
    2a. LlamaIndex SemanticSplitterNodeParser (Section 4.2.1)
        - Embedding-based splitting at topic shifts
    2b. Hierarchical Clustering (Section 4.2.2) - Enhanced option
        - Multi-resolution: micro-clusters → macro-clusters
    2c. Synthetic Header Generation (Section 6.3)
        - LLM-generated titles for each section

TIER 3: OCR + Re-analyze (Fallback 2 - Scanned PDFs)
    - Apply Tesseract or Doctr OCR
    - Generate text layer from images
    - Build tree from OCR output

ALWAYS call `ingest_document()` - never call individual tiers directly.
"""

from __future__ import annotations

from pathlib import Path

import structlog

from rnsr.exceptions import IngestionError
from rnsr.ingestion.document_boundary import (
    DocumentBoundaryDetector,
    segment_by_documents,
)
from rnsr.ingestion.font_histogram import FontHistogramAnalyzer
from rnsr.ingestion.header_classifier import classify_headers
from rnsr.ingestion.layout_detector import detect_layout_complexity
from rnsr.ingestion.ocr_fallback import has_extractable_text, try_ocr_ingestion
from rnsr.ingestion.semantic_fallback import try_semantic_splitter_ingestion
from rnsr.ingestion.tree_builder import build_document_tree, build_multi_document_tree
from rnsr.models import DetectedTable, DocumentNode, DocumentTree, IngestionResult

logger = structlog.get_logger(__name__)


def _extract_docx_text(docx_path: Path) -> str | None:
    """Extract text from a .docx file including paragraphs, tables, headers,
    footers, and structured document tags (content controls / form fields).

    Walks the raw XML so nothing is missed -- python-docx helpers like
    ``doc.paragraphs`` skip text inside ``<w:sdt>`` tags and don't expose
    tables that live inside headers/footers.
    """
    try:
        import docx  # python-docx
    except ImportError:
        logger.warning("python_docx_not_installed", hint="pip install python-docx")
        return None

    _WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    _W_P = f"{{{_WNS}}}p"
    _W_R = f"{{{_WNS}}}r"
    _W_TBL = f"{{{_WNS}}}tbl"
    _W_SDT = f"{{{_WNS}}}sdt"
    _W_SDT_CONTENT = f"{{{_WNS}}}sdtContent"
    _W_TR = f"{{{_WNS}}}tr"
    _W_TC = f"{{{_WNS}}}tc"
    _W_T = f"{{{_WNS}}}t"
    _W_BR = f"{{{_WNS}}}br"
    _W_CR = f"{{{_WNS}}}cr"
    _W_TAB = f"{{{_WNS}}}tab"

    def _para_text(p_elem) -> str:
        """Collect all <w:t> text from a paragraph element, respecting
        breaks, tabs and content controls nested inline."""
        fragments: list[str] = []
        for node in p_elem.iter():
            if node.tag == _W_T and node.text:
                fragments.append(node.text)
            elif node.tag in (_W_BR, _W_CR):
                fragments.append("\n")
            elif node.tag == _W_TAB:
                fragments.append("\t")
        return "".join(fragments).strip()

    def _table_text(tbl_elem) -> list[str]:
        """Extract table rows as pipe-separated cell text to preserve
        structure.  Handles SDTs and nested tables inside cells."""
        rows: list[str] = []
        for tr in tbl_elem:
            if tr.tag != _W_TR:
                continue
            cells: list[str] = []
            for tc in tr:
                if tc.tag != _W_TC:
                    continue
                cell_parts: list[str] = []
                for child in tc:
                    if child.tag == _W_P:
                        t = _para_text(child)
                        if t:
                            cell_parts.append(t)
                    elif child.tag == _W_TBL:
                        cell_parts.extend(_table_text(child))
                    elif child.tag == _W_SDT:
                        cell_parts.extend(_process_sdt(child))
                cell_text = " ".join(cell_parts).strip()
                if cell_text:
                    cells.append(cell_text)
            if cells:
                rows.append(" | ".join(cells))
        return rows

    def _run_text(r_elem) -> str:
        """Collect text from a bare <w:r> element."""
        fragments: list[str] = []
        for node in r_elem.iter():
            if node.tag == _W_T and node.text:
                fragments.append(node.text)
            elif node.tag in (_W_BR, _W_CR):
                fragments.append("\n")
            elif node.tag == _W_TAB:
                fragments.append("\t")
        return "".join(fragments).strip()

    def _process_sdt(sdt_elem) -> list[str]:
        """Recurse into an SDT's content -- it can hold paragraphs, tables,
        bare runs, or further nested SDTs."""
        parts: list[str] = []
        for child in sdt_elem:
            if child.tag == _W_SDT_CONTENT:
                for inner in child:
                    if inner.tag == _W_P:
                        t = _para_text(inner)
                        if t:
                            parts.append(t)
                    elif inner.tag == _W_TBL:
                        parts.extend(_table_text(inner))
                    elif inner.tag == _W_SDT:
                        parts.extend(_process_sdt(inner))
                    elif inner.tag == _W_R:
                        t = _run_text(inner)
                        if t:
                            parts.append(t)
        return parts

    def _extract_container(container_elem) -> list[str]:
        """Walk direct children of a container element (body, header,
        footer) and extract text in document order without duplication."""
        parts: list[str] = []
        for child in container_elem:
            if child.tag == _W_P:
                t = _para_text(child)
                if t:
                    parts.append(t)
            elif child.tag == _W_TBL:
                parts.extend(_table_text(child))
            elif child.tag == _W_SDT:
                parts.extend(_process_sdt(child))
        return parts

    try:
        doc = docx.Document(str(docx_path))
        parts: list[str] = []

        for section in doc.sections:
            for hf in (section.header, section.footer):
                if hf and hf._element is not None:
                    parts.extend(_extract_container(hf._element))

        parts.extend(_extract_container(doc.element.body))

        return "\n\n".join(parts) if parts else None
    except Exception as exc:
        logger.warning("docx_extraction_failed", path=str(docx_path), error=str(exc))
        return None


def _extract_xlsx_text(xlsx_path: Path) -> str | None:
    """Extract text from an Excel workbook (.xlsx / .xls).

    Each sheet is rendered as a section with its name as header, followed by
    rows formatted as pipe-separated values so table structure is preserved.
    """
    try:
        import openpyxl
    except ImportError:
        logger.warning("openpyxl_not_installed", hint="pip install openpyxl")
        return None

    try:
        wb = openpyxl.load_workbook(str(xlsx_path), read_only=True, data_only=True)
        parts: list[str] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_rows: list[str] = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() if c is not None else "" for c in row]
                if any(cells):
                    sheet_rows.append(" | ".join(cells))
            if sheet_rows:
                parts.append(f"[Sheet: {sheet_name}]")
                parts.extend(sheet_rows)

        wb.close()
        return "\n\n".join(parts) if parts else None
    except Exception as exc:
        logger.warning("xlsx_extraction_failed", path=str(xlsx_path), error=str(exc))
        return None


def _extract_xlsx_tables(xlsx_path: Path) -> list:
    """Build DetectedTable objects directly from Excel workbook data.

    Each non-empty worksheet becomes a DetectedTable with the first non-empty
    row treated as column headers and subsequent rows as data.  This bypasses
    the text-based table parser which cannot detect tables from the ``\\n\\n``
    separated text produced by ``_extract_xlsx_text()``.
    """
    from rnsr.models import DetectedTable

    try:
        import openpyxl
    except ImportError:
        logger.warning("openpyxl_not_installed", hint="pip install openpyxl")
        return []

    try:
        import hashlib

        wb = openpyxl.load_workbook(str(xlsx_path), read_only=True, data_only=True)
        tables: list[DetectedTable] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            all_rows: list[list[str]] = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() if c is not None else "" for c in row]
                if any(cells):
                    all_rows.append(cells)

            if len(all_rows) < 2:
                continue

            headers = all_rows[0]
            data_rows = all_rows[1:]

            table_id = hashlib.sha256(
                f"{xlsx_path.name}:{sheet_name}".encode()
            ).hexdigest()[:12]

            tables.append(
                DetectedTable(
                    id=table_id,
                    node_id="root",
                    title=sheet_name,
                    headers=headers,
                    num_rows=len(data_rows),
                    num_cols=len(headers),
                    data=data_rows,
                )
            )

        wb.close()
        logger.info("xlsx_tables_built", path=str(xlsx_path), count=len(tables))
        return tables
    except Exception as exc:
        logger.warning("xlsx_table_extraction_failed", path=str(xlsx_path), error=str(exc))
        return []


def _extract_msg_text(msg_path: Path) -> str | None:
    """Extract text from an Outlook .msg email file.

    Returns structured text with From/To/Subject/Date headers and body.
    """
    try:
        import extract_msg
    except ImportError:
        logger.warning("extract_msg_not_installed", hint="pip install extract-msg")
        return None

    try:
        msg = extract_msg.openMsg(str(msg_path))
        parts: list[str] = []
        if msg.subject:
            parts.append(f"Subject: {msg.subject}")
        if msg.sender:
            parts.append(f"From: {msg.sender}")
        if msg.to:
            parts.append(f"To: {msg.to}")
        if msg.date:
            parts.append(f"Date: {msg.date}")
        if msg.cc:
            parts.append(f"CC: {msg.cc}")
        if parts:
            parts.append("")  # blank line before body
        if msg.body:
            parts.append(msg.body.strip())
        msg.close()
        return "\n".join(parts) if parts else None
    except Exception as exc:
        logger.warning("msg_extraction_failed", path=str(msg_path), error=str(exc))
        return None


def _extract_csv_text(csv_path: Path) -> str | None:
    """Extract text from a CSV file, rendering rows as pipe-separated values."""
    try:
        import csv

        with open(csv_path, newline="", encoding="utf-8-sig") as f:
            reader = csv.reader(f)
            rows: list[str] = []
            for row in reader:
                cells = [c.strip() for c in row]
                if any(cells):
                    rows.append(" | ".join(cells))
        return "\n\n".join(rows) if rows else None
    except Exception as exc:
        logger.warning("csv_extraction_failed", path=str(csv_path), error=str(exc))
        return None


def _extract_csv_tables(csv_path: Path) -> list:
    """Build a single DetectedTable from a CSV file.

    The first non-empty row is treated as column headers.
    """
    from rnsr.models import DetectedTable

    try:
        import csv
        import hashlib

        with open(csv_path, newline="", encoding="utf-8-sig") as f:
            reader = csv.reader(f)
            all_rows: list[list[str]] = []
            for row in reader:
                cells = [c.strip() for c in row]
                if any(cells):
                    all_rows.append(cells)

        if len(all_rows) < 2:
            return []

        headers = all_rows[0]
        data_rows = all_rows[1:]

        table_id = hashlib.sha256(csv_path.name.encode()).hexdigest()[:12]

        table = DetectedTable(
            id=table_id,
            node_id="root",
            title=csv_path.stem,
            headers=headers,
            num_rows=len(data_rows),
            num_cols=len(headers),
            data=data_rows,
        )
        logger.info("csv_table_built", path=str(csv_path), rows=len(data_rows))
        return [table]
    except Exception as exc:
        logger.warning("csv_table_extraction_failed", path=str(csv_path), error=str(exc))
        return []


_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff", ".tif"}


def _extract_image_text(image_path: Path) -> str | None:
    """Extract text from an image using a Vision Language Model (VLM).

    Uses the configured Gemini model to describe and transcribe all visible
    text in the image (screenshots, scanned documents, photos of text, etc.).
    """
    try:
        from rnsr.llm import get_llm, LLMProvider
    except ImportError:
        logger.warning("llm_module_not_available_for_vlm")
        return None

    try:
        image_bytes = image_path.read_bytes()

        import mimetypes
        mime, _ = mimetypes.guess_type(str(image_path))
        if not mime:
            mime = "image/png"

        llm = get_llm(provider=LLMProvider.GEMINI, enable_fallback=True)

        prompt = (
            "You are a document extraction assistant. Transcribe ALL visible text "
            "in this image exactly as it appears, preserving layout, line breaks, "
            "and any structure (tables, lists, headings). If the image is a "
            "screenshot of a conversation or text message, transcribe each message "
            "with its sender. Output ONLY the transcribed text, nothing else."
        )

        result = llm.complete_with_image(prompt, image_bytes)
        text = str(result).strip()
        if text:
            logger.info("vlm_image_extracted", path=str(image_path), chars=len(text))
            return text
        return None
    except Exception as exc:
        logger.warning("image_extraction_failed", path=str(image_path), error=str(exc))
        return None


def _text_to_ingestion_result(
    text: str,
    file_path: Path,
    stats: dict,
    warnings: list[str],
) -> IngestionResult:
    """Convert extracted plain text into an IngestionResult with optional
    pattern-based header detection."""
    from rnsr.ingestion.semantic_fallback import _try_pattern_based_headers

    tree = _try_pattern_based_headers(text, file_path.stem)
    if tree:
        result = IngestionResult(
            tree=tree,
            tier_used=2,
            method="pattern_based_headers",
            stats=stats,
            warnings=warnings,
        )
        return _add_tables_to_result(result)

    root = DocumentNode(id="root", level=0, header=file_path.stem, content=text)
    tree = DocumentTree(
        title=file_path.stem,
        root=root,
        total_nodes=1,
        ingestion_tier=2,
        ingestion_method="semantic_splitter",
    )
    result = IngestionResult(
        tree=tree,
        tier_used=2,
        method="semantic_splitter",
        stats=stats,
        warnings=warnings,
    )
    return _add_tables_to_result(result)


def ingest_document(
    pdf_path: Path | str,
    use_visual_analysis: bool = True,
    complexity_threshold: float = 0.3,
) -> IngestionResult:
    """
    Master ingestion function implementing 3-tier graceful degradation.
    
    ALWAYS call this function - never call individual tiers directly.
    
    Ingestion Flow:
    0. Pre-analysis: Detect layout complexity (multi-column, empty pages)
    1. Tier 1a: Font Histogram (simple layouts)
    1. Tier 1b: LayoutLM + XY-Cut (complex layouts, if use_visual_analysis=True)
    2. Tier 2: Semantic Splitter (flat text, no structure)
    3. Tier 3: OCR (scanned/image-only PDFs)
    
    Args:
        pdf_path: Path to the PDF file to ingest.
        use_visual_analysis: Enable LayoutLM for complex layouts (default: True).
        complexity_threshold: Threshold for triggering visual analysis (0.0-1.0).
        
    Returns:
        IngestionResult containing the DocumentTree and metadata.
        
    Raises:
        IngestionError: If all tiers fail.
        
    Example:
        # Auto-detect layout complexity
        result = ingest_document("contract.pdf")
        
        # Force visual analysis
        result = ingest_document("report.pdf", use_visual_analysis=True)
        
        # Disable visual analysis
        result = ingest_document("simple.pdf", use_visual_analysis=False)
    """
    pdf_path = Path(pdf_path)
    
    if not pdf_path.exists():
        raise IngestionError(f"Document file not found: {pdf_path}")
    
    logger.info("ingestion_started", path=str(pdf_path))
    
    warnings: list[str] = []
    stats: dict = {"path": str(pdf_path)}
    
    suffix = pdf_path.suffix.lower()

    # --- Non-PDF file types: extract text then build tree -----------------

    if suffix in {".md", ".txt", ".text", ".markdown"}:
        logger.info("text_file_detected", path=str(pdf_path))
        text = pdf_path.read_text(encoding="utf-8")
        return _text_to_ingestion_result(text, pdf_path, stats, warnings)

    if suffix == ".docx":
        logger.info("docx_file_detected", path=str(pdf_path))
        text = _extract_docx_text(pdf_path)
        if not text:
            raise IngestionError(f"Failed to extract text from docx: {pdf_path}")
        return _text_to_ingestion_result(text, pdf_path, stats, warnings)

    if suffix in {".xlsx", ".xls"}:
        logger.info("xlsx_file_detected", path=str(pdf_path))
        text = _extract_xlsx_text(pdf_path)
        if not text:
            raise IngestionError(f"Failed to extract text from Excel: {pdf_path}")
        result = _text_to_ingestion_result(text, pdf_path, stats, warnings)
        excel_tables = _extract_xlsx_tables(pdf_path)
        if excel_tables:
            result.tables = excel_tables
            result.stats["tables_detected"] = len(excel_tables)
        return result

    if suffix == ".msg":
        logger.info("msg_file_detected", path=str(pdf_path))
        text = _extract_msg_text(pdf_path)
        if not text:
            raise IngestionError(f"Failed to extract text from .msg: {pdf_path}")
        return _text_to_ingestion_result(text, pdf_path, stats, warnings)

    if suffix == ".csv":
        logger.info("csv_file_detected", path=str(pdf_path))
        text = _extract_csv_text(pdf_path)
        if not text:
            raise IngestionError(f"Failed to extract text from CSV: {pdf_path}")
        result = _text_to_ingestion_result(text, pdf_path, stats, warnings)
        csv_tables = _extract_csv_tables(pdf_path)
        if csv_tables:
            result.tables = csv_tables
            result.stats["tables_detected"] = len(csv_tables)
        return result

    if suffix in _IMAGE_EXTENSIONS:
        logger.info("image_file_detected", path=str(pdf_path))
        text = _extract_image_text(pdf_path)
        if not text:
            raise IngestionError(
                f"Failed to extract text from image via VLM: {pdf_path}"
            )
        return _text_to_ingestion_result(text, pdf_path, stats, warnings)

    # --- PDF path (existing tier logic) -----------------------------------

    # Check if document has extractable text
    if not has_extractable_text(pdf_path):
        # No text - go directly to Tier 3 (OCR)
        logger.info("no_extractable_text", path=str(pdf_path))
        return _try_tier_3(pdf_path, warnings, stats)
    
    # PRE-ANALYSIS: Detect layout complexity
    if use_visual_analysis:
        try:
            complexity = detect_layout_complexity(pdf_path, threshold=complexity_threshold)
            
            stats["layout_complexity"] = complexity.complexity_score
            stats["needs_visual"] = complexity.needs_visual_analysis
            stats["complexity_reason"] = complexity.reason
            
            if complexity.needs_visual_analysis:
                logger.info(
                    "complex_layout_detected",
                    path=str(pdf_path),
                    score=complexity.complexity_score,
                    reason=complexity.reason,
                )
                
                # Try visual analysis first
                result = _try_tier_1b_visual(pdf_path, warnings, stats)
                if result is not None:
                    return result
                
                # Fall through to standard font histogram if visual fails
                warnings.append(f"Visual analysis failed, using font histogram fallback")
        except Exception as e:
            logger.warning("layout_detection_failed", error=str(e))
            warnings.append(f"Layout detection failed: {e}")
    
    # TIER 1: Try PyMuPDF Font Histogram
    result = _try_tier_1(pdf_path, warnings, stats)
    if result is not None:
        return result
    
    # TIER 2: Try Semantic Splitter
    result = _try_tier_2(pdf_path, warnings, stats)
    if result is not None:
        return result
    
    # This shouldn't happen, but just in case
    raise IngestionError("All ingestion tiers failed")


def _try_tier_1b_visual(
    pdf_path: Path,
    warnings: list[str],
    stats: dict,
) -> IngestionResult | None:
    """
    TIER 1b: Try LayoutLM + XY-Cut for complex layouts.
    
    Uses visual analysis to detect document structure when
    layout is too complex for simple font histogram.
    """
    logger.debug("trying_tier_1b_visual", path=str(pdf_path))
    
    try:
        from rnsr.ingestion.layout_model import check_layout_model_available
        
        if not check_layout_model_available():
            logger.warning("layout_model_unavailable")
            warnings.append("LayoutLM not available - falling back to font histogram")
            return None
        
        from rnsr.ingestion.xy_cut import analyze_document_with_xycut
        
        # Use XY-Cut + LayoutLM for visual analysis
        tree = analyze_document_with_xycut(pdf_path)
        tree.ingestion_tier = 1
        tree.ingestion_method = "layoutlm_xycut"
        
        logger.info(
            "tier_1b_visual_success",
            path=str(pdf_path),
            nodes=tree.total_nodes,
        )
        
        result = IngestionResult(
            tree=tree,
            tier_used=1,
            method="layoutlm_xycut",
            warnings=warnings,
            stats=stats,
        )
        return _add_tables_to_result(result)
        
    except Exception as e:
        logger.warning("tier_1b_visual_failed", path=str(pdf_path), error=str(e))
        warnings.append(f"LayoutLM visual analysis failed: {e}")
        return None


def _try_tier_1(
    pdf_path: Path,
    warnings: list[str],
    stats: dict,
    detect_multi_document: bool = True,
    boundary_confidence: float = 0.5,
) -> IngestionResult | None:
    """
    TIER 1: Try Font Histogram ingestion.
    
    Now includes multi-document detection for combined PDFs.
    
    Returns None if should fall back to Tier 2.
    """
    logger.debug("trying_tier_1", path=str(pdf_path))
    
    try:
        analyzer = FontHistogramAnalyzer()
        analysis, spans = analyzer.analyze(pdf_path)
        
        stats["span_count"] = len(spans)
        stats["unique_sizes"] = analysis.unique_sizes
        stats["body_size"] = analysis.body_size
        
        # Check if we have font variance
        if not analyzer.has_font_variance(analysis):
            logger.info("no_font_variance", path=str(pdf_path))
            warnings.append("No font variance detected - using semantic splitter")
            return None  # Trigger Tier 2
        
        # Check if we can detect headers
        if not analyzer.has_detectable_headers(analysis, spans):
            logger.info("no_headers_detected", path=str(pdf_path))
            warnings.append("No headers detected - using semantic splitter")
            return None  # Trigger Tier 2
        
        # NEW: Detect document boundaries for multi-document PDFs
        if detect_multi_document:
            segments = segment_by_documents(
                spans, 
                min_confidence=boundary_confidence,
            )
            
            stats["documents_detected"] = len(segments)
            
            if len(segments) > 1:
                logger.info(
                    "multi_document_detected",
                    path=str(pdf_path),
                    document_count=len(segments),
                    titles=[s.title[:30] for s in segments],
                )
                
                # Build multi-document tree
                tree = build_multi_document_tree(
                    segments,
                    container_title=pdf_path.stem,
                )
                tree.ingestion_tier = 1
                tree.ingestion_method = "font_histogram"
                
                logger.info(
                    "tier_1_success",
                    path=str(pdf_path),
                    nodes=tree.total_nodes,
                    documents=len(segments),
                )
                
                result = IngestionResult(
                    tree=tree,
                    tier_used=1,
                    method="font_histogram",
                    warnings=warnings,
                    stats=stats,
                )
                return _add_tables_to_result(result)
        
        # Single document: standard processing
        # Classify spans
        classified = classify_headers(spans, analysis)
        
        header_count = sum(1 for s in classified if s.role == "header")
        stats["header_count"] = header_count
        
        # Build tree
        tree = build_document_tree(classified, title=pdf_path.stem)
        tree.ingestion_tier = 1
        tree.ingestion_method = "font_histogram"
        
        logger.info(
            "tier_1_success",
            path=str(pdf_path),
            nodes=tree.total_nodes,
        )
        
        result = IngestionResult(
            tree=tree,
            tier_used=1,
            method="font_histogram",
            warnings=warnings,
            stats=stats,
        )
        return _add_tables_to_result(result)
        
    except Exception as e:
        logger.warning("tier_1_failed", path=str(pdf_path), error=str(e))
        warnings.append(f"Font histogram failed: {e}")
        return None


def _try_tier_2(
    pdf_path: Path,
    warnings: list[str],
    stats: dict,
    use_hierarchical_clustering: bool = False,
) -> IngestionResult | None:
    """
    TIER 2: Try Semantic Splitter or Hierarchical Clustering ingestion.
    
    Implements Section 4.2 of the research paper:
    - 4.2.1: SemanticSplitterNodeParser for breakpoint detection
    - 4.2.2: Hierarchical Clustering for multi-resolution topics
    - 6.3: Synthetic Header Generation via LLM
    """
    logger.debug("trying_tier_2", path=str(pdf_path))
    
    # Option: Use hierarchical clustering for richer structure
    if use_hierarchical_clustering:
        try:
            from rnsr.ingestion.hierarchical_cluster import cluster_document_hierarchically
            
            tree = cluster_document_hierarchically(pdf_path)
            
            logger.info(
                "tier_2_hierarchical_success",
                path=str(pdf_path),
                nodes=tree.total_nodes,
            )
            
            result = IngestionResult(
                tree=tree,
                tier_used=2,
                method="hierarchical_clustering",
                warnings=warnings,
                stats=stats,
            )
            return _add_tables_to_result(result)
        except Exception as e:
            logger.warning("hierarchical_clustering_failed", error=str(e))
            warnings.append(f"Hierarchical clustering failed: {e}")
            # Fall through to semantic splitter
    
    # Default: Semantic Splitter (with LLM-generated headers)
    try:
        tree = try_semantic_splitter_ingestion(pdf_path)
        
        logger.info(
            "tier_2_success",
            path=str(pdf_path),
            nodes=tree.total_nodes,
        )
        
        result = IngestionResult(
            tree=tree,
            tier_used=2,
            method="semantic_splitter",
            warnings=warnings,
            stats=stats,
        )
        return _add_tables_to_result(result)
        
    except Exception as e:
        logger.warning("tier_2_failed", path=str(pdf_path), error=str(e))
        warnings.append(f"Semantic splitter failed: {e}")
        # Continue to Tier 3
        return _try_tier_3(pdf_path, warnings, stats)


def _try_tier_3(
    pdf_path: Path,
    warnings: list[str],
    stats: dict,
) -> IngestionResult:
    """
    TIER 3: Try OCR ingestion (last resort).
    """
    logger.debug("trying_tier_3", path=str(pdf_path))
    
    try:
        tree = try_ocr_ingestion(pdf_path)
        
        logger.info(
            "tier_3_success",
            path=str(pdf_path),
            nodes=tree.total_nodes,
        )
        
        result = IngestionResult(
            tree=tree,
            tier_used=3,
            method=tree.ingestion_method or "ocr",
            warnings=warnings,
            stats=stats,
        )
        return _add_tables_to_result(result)
        
    except Exception as e:
        logger.error("tier_3_failed", path=str(pdf_path), error=str(e))
        raise IngestionError(f"All ingestion tiers failed. Last error: {e}") from e


def ingest_document_enhanced(
    pdf_path: Path | str,
    use_xy_cut: bool = False,
    use_hierarchical_clustering: bool = False,
) -> IngestionResult:
    """
    Enhanced ingestion with all research paper features.
    
    This exposes the full Latent Hierarchy Generator from the paper:
    - XY-Cut for complex multi-column layouts (Section 4.1.1)
    - Hierarchical Clustering for multi-resolution topics (Section 4.2.2)
    - Synthetic Header Generation via LLM (Section 6.3)
    
    Args:
        pdf_path: Path to the PDF file to ingest.
        use_xy_cut: Enable Recursive XY-Cut for complex layouts.
        use_hierarchical_clustering: Use clustering instead of simple splits.
        
    Returns:
        IngestionResult containing the DocumentTree and metadata.
        
    Example:
        # For a complex multi-column PDF:
        result = ingest_document_enhanced("report.pdf", use_xy_cut=True)
        
        # For flat text that needs hierarchical structure:
        result = ingest_document_enhanced(
            "transcript.pdf",
            use_hierarchical_clustering=True
        )
    """
    pdf_path = Path(pdf_path)
    
    if not pdf_path.exists():
        raise IngestionError(f"PDF file not found: {pdf_path}")
    
    logger.info(
        "enhanced_ingestion_started",
        path=str(pdf_path),
        xy_cut=use_xy_cut,
        hierarchical=use_hierarchical_clustering,
    )
    
    warnings: list[str] = []
    stats: dict = {"path": str(pdf_path)}
    
    # Check if document has extractable text
    if not has_extractable_text(pdf_path):
        return _try_tier_3(pdf_path, warnings, stats)
    
    # Try XY-Cut first if enabled (for complex layouts)
    if use_xy_cut:
        result = _try_xy_cut_ingestion(pdf_path, warnings, stats)
        if result is not None:
            return result
    
    # TIER 1: Try Font Histogram
    result = _try_tier_1(pdf_path, warnings, stats)
    if result is not None:
        return result
    
    # TIER 2: Semantic analysis with optional hierarchical clustering
    result = _try_tier_2(pdf_path, warnings, stats, use_hierarchical_clustering)
    if result is not None:
        return result
    
    raise IngestionError("All ingestion tiers failed")


def _try_xy_cut_ingestion(
    pdf_path: Path,
    warnings: list[str],
    stats: dict,
) -> IngestionResult | None:
    """
    Optional: Use Recursive XY-Cut + LayoutLM for complex layouts.
    
    Implements Section 4.1.1:
    "A top-down page segmentation technique that is particularly 
    effective for discovering document structure."
    """
    logger.debug("trying_xy_cut_with_layoutlm", path=str(pdf_path))
    
    try:
        # Check if LayoutLM is available
        from rnsr.ingestion.layout_model import check_layout_model_available
        
        if not check_layout_model_available():
            logger.warning("layout_model_unavailable")
            warnings.append("LayoutLM not available for XY-Cut enhancement")
            return None
        
        from rnsr.ingestion.xy_cut import analyze_document_with_xycut
        
        # Use XY-Cut + LayoutLM for visual analysis
        tree = analyze_document_with_xycut(pdf_path)
        tree.ingestion_tier = 1
        tree.ingestion_method = "xy_cut_layoutlm"
        
        logger.info(
            "xy_cut_layoutlm_success",
            path=str(pdf_path),
            nodes=tree.total_nodes,
        )
        
        result = IngestionResult(
            tree=tree,
            tier_used=1,
            method="xy_cut_layoutlm",
            warnings=warnings,
            stats=stats,
        )
        return _add_tables_to_result(result)
        
    except Exception as e:
        logger.warning("xy_cut_layoutlm_failed", path=str(pdf_path), error=str(e))
        warnings.append(f"XY-Cut + LayoutLM failed: {e}")
        return None


def _try_xy_cut_ingestion_legacy(
    pdf_path: Path,
    warnings: list[str],
    stats: dict,
) -> IngestionResult | None:
    """
    Legacy XY-Cut implementation without LayoutLM.
    
    Implements Section 4.1.1:
    "A top-down page segmentation technique that is particularly 
    effective for discovering document structure."
    """
    logger.debug("trying_xy_cut", path=str(pdf_path))
    
    try:
        from rnsr.ingestion.xy_cut import RecursiveXYCutter
        import fitz
        
        cutter = RecursiveXYCutter()
        page_trees = cutter.segment_pdf(pdf_path)
        
        # Extract text for each leaf region
        doc = fitz.open(pdf_path)
        for page_num, tree in enumerate(page_trees):
            cutter.extract_text_in_regions(doc[page_num], tree)
        doc.close()
        
        # Convert XY-Cut tree to DocumentTree
        from rnsr.models import DocumentNode, DocumentTree
        
        root = DocumentNode(
            id="root",
            level=0,
            header=pdf_path.stem,
        )
        
        section_num = 0
        for page_tree in page_trees:
            for leaf in _get_xy_cut_leaves(page_tree):
                if leaf.text.strip():
                    section_num += 1
                    # Generate synthetic header
                    from rnsr.ingestion.semantic_fallback import _generate_synthetic_header
                    
                    section = DocumentNode(
                        id=f"xycut_{section_num:03d}",
                        level=1,
                        header=_generate_synthetic_header(leaf.text, section_num),
                        content=leaf.text,
                    )
                    root.children.append(section)
        
        if section_num == 0:
            warnings.append("XY-Cut found no text regions")
            return None
        
        tree = DocumentTree(
            title=pdf_path.stem,
            root=root,
            total_nodes=section_num + 1,
            ingestion_tier=1,
            ingestion_method="xy_cut",
        )
        
        logger.info("xy_cut_success", path=str(pdf_path), nodes=tree.total_nodes)
        
        result = IngestionResult(
            tree=tree,
            tier_used=1,
            method="xy_cut",
            warnings=warnings,
            stats=stats,
        )
        return _add_tables_to_result(result)
        
    except Exception as e:
        logger.warning("xy_cut_failed", path=str(pdf_path), error=str(e))
        warnings.append(f"XY-Cut failed: {e}")
        return None


def _get_xy_cut_leaves(node) -> list:
    """Get all leaf nodes from an XY-Cut segment tree."""
    if node.is_leaf:
        return [node]
    leaves = []
    for child in node.children:
        leaves.extend(_get_xy_cut_leaves(child))
    return leaves


# =============================================================================
# Entity Extraction Integration
# =============================================================================


def extract_entities_from_tree(
    tree: DocumentTree,
    doc_id: str | None = None,
    extract_relationships: bool = True,
    max_nodes: int | None = None,
    sample_strategy: str = "all",
) -> dict:
    """
    Extract entities and relationships from an ingested document tree.
    
    Uses the RLM Unified Extractor - the most accurate approach:
    1. LLM writes extraction code based on document
    2. Code executes on DOC_VAR (grounded in actual text)
    3. ToT validates with probabilities
    4. Cross-validation between entities and relationships
    
    Args:
        tree: The ingested DocumentTree.
        doc_id: Document ID (defaults to tree.id).
        extract_relationships: Whether to also extract relationships.
        max_nodes: Optional cap on nodes to process.  Defaults to ``None``
            (process every node).  Set to a positive integer to limit
            extraction for very large documents.
        sample_strategy: How to select nodes when *max_nodes* triggers —
            ``"all"`` (default, process every node), ``"important"``
            (headers first), or ``"uniform"`` (evenly spaced).
        
    Returns:
        Dictionary containing:
        - entities: List of extracted Entity objects
        - relationships: List of extracted Relationship objects
        - stats: Extraction statistics
        
    Example:
        result = ingest_document("contract.pdf")
        extraction = extract_entities_from_tree(result.tree)
        
        # Store in knowledge graph
        for entity in extraction["entities"]:
            kg.add_entity(entity)
    """
    from rnsr.extraction import (
        RLMUnifiedExtractor,
        merge_entities,
    )
    
    doc_id = doc_id or tree.id
    
    # Use RLM Unified Extractor (LLM writes code + ToT validation)
    extractor = RLMUnifiedExtractor(
        enable_type_learning=True,
        enable_tot_validation=True,
        enable_cross_validation=True,
    )
    
    all_entities = []
    all_relationships = []
    
    # Collect all nodes for processing
    all_nodes = _collect_nodes(tree.root, doc_id)
    
    # Sample nodes if a cap is set and exceeded
    if max_nodes is not None and sample_strategy != "all" and len(all_nodes) > max_nodes:
        nodes_to_process = _sample_nodes(all_nodes, max_nodes, sample_strategy)
        logger.info(
            "entity_extraction_sampling",
            total_nodes=len(all_nodes),
            sampled_nodes=len(nodes_to_process),
            strategy=sample_strategy,
        )
    else:
        nodes_to_process = all_nodes
    
    logger.info(
        "entity_extraction_started",
        doc_id=doc_id,
        node_count=len(nodes_to_process),
        total_nodes=len(all_nodes),
        extractor="RLMUnifiedExtractor",
    )
    
    # Process nodes in batches for efficiency
    batch_size = 10
    processed = 0
    
    for i in range(0, len(nodes_to_process), batch_size):
        batch = nodes_to_process[i:i + batch_size]
        
        for node_data in batch:
            try:
                result = extractor.extract(
                    node_id=node_data["node_id"],
                    doc_id=doc_id,
                    header=node_data["header"],
                    content=node_data["content"],
                    page_num=node_data.get("page_num"),
                )
                
                if result.entities:
                    all_entities.extend(result.entities)
                
                if extract_relationships and result.relationships:
                    all_relationships.extend(result.relationships)
                    
            except Exception as e:
                logger.warning(
                    "node_extraction_failed",
                    node_id=node_data.get("node_id"),
                    error=str(e)[:100],
                )
        
        processed += len(batch)
        if processed % 20 == 0:
            logger.info(
                "entity_extraction_progress",
                processed=processed,
                total=len(nodes_to_process),
                entities_so_far=len(all_entities),
            )
    
    # Merge duplicate entities
    merged_entities = merge_entities(all_entities)
    
    stats = {
        "nodes_processed": len(nodes_to_process),
        "entities_extracted": len(all_entities),
        "entities_after_merge": len(merged_entities),
        "relationships_extracted": len(all_relationships),
        "entity_types": _count_entity_types(merged_entities),
        "extraction_method": "rlm_unified",
    }
    
    logger.info(
        "entity_extraction_complete",
        doc_id=doc_id,
        **stats,
    )
    
    return {
        "entities": merged_entities,
        "relationships": all_relationships,
        "stats": stats,
    }


def _collect_nodes(node, doc_id: str, collected: list | None = None) -> list[dict]:
    """
    Recursively collect all nodes from a DocumentNode tree.
    
    Args:
        node: Root DocumentNode.
        doc_id: Document ID.
        collected: List to collect into.
        
    Returns:
        List of node data dictionaries.
    """
    if collected is None:
        collected = []
    
    # Add this node if it has content
    if node.content or node.header:
        collected.append({
            "node_id": node.id,
            "header": node.header,
            "content": node.content,
            "page_num": node.page_num,
            "level": node.level,
        })
    
    # Process children
    for child in node.children:
        _collect_nodes(child, doc_id, collected)
    
    return collected


def _sample_nodes(nodes: list[dict], max_nodes: int, strategy: str) -> list[dict]:
    """
    Sample nodes from a large document for efficient processing.
    
    Args:
        nodes: All collected nodes.
        max_nodes: Maximum number of nodes to return.
        strategy: Sampling strategy - "important", "uniform", or "first".
        
    Returns:
        Sampled list of nodes.
    """
    if len(nodes) <= max_nodes:
        return nodes
    
    if strategy == "important":
        # Prioritize nodes with headers and higher-level sections
        scored_nodes = []
        for node in nodes:
            score = 0
            
            # Prefer nodes with headers
            if node.get("header"):
                score += 10
            
            # Prefer higher-level (lower number) sections
            level = node.get("level", 3)
            score += max(0, 5 - level)
            
            # Prefer nodes with substantial content
            content_len = len(node.get("content", ""))
            if content_len > 500:
                score += 3
            elif content_len > 200:
                score += 2
            elif content_len > 50:
                score += 1
            
            scored_nodes.append((score, node))
        
        # Sort by score (descending) and take top nodes
        scored_nodes.sort(key=lambda x: x[0], reverse=True)
        return [node for _, node in scored_nodes[:max_nodes]]
    
    elif strategy == "uniform":
        # Evenly sample across the document
        step = len(nodes) // max_nodes
        return [nodes[i] for i in range(0, len(nodes), step)][:max_nodes]
    
    else:  # "first"
        return nodes[:max_nodes]


def _count_entity_types(entities: list) -> dict[str, int]:
    """Count entities by type."""
    counts: dict[str, int] = {}
    for entity in entities:
        type_name = entity.type.value
        counts[type_name] = counts.get(type_name, 0) + 1
    return counts


def ingest_with_entities(
    pdf_path: Path | str,
    knowledge_graph=None,
    extract_relationships: bool = True,
    link_entities: bool = True,
    **ingest_kwargs,
) -> dict:
    """
    Ingest a document and extract entities in a single operation.
    
    Combines document ingestion with entity extraction, optionally
    storing results in a knowledge graph.
    
    Args:
        pdf_path: Path to the PDF file.
        knowledge_graph: Optional KnowledgeGraph to store entities.
        extract_relationships: Whether to extract relationships.
        link_entities: Whether to link entities across documents.
        **ingest_kwargs: Additional arguments for ingest_document.
        
    Returns:
        Dictionary containing:
        - ingestion_result: The IngestionResult
        - extraction: Entity extraction results
        - links: Entity links (if link_entities=True)
        
    Example:
        from rnsr.indexing.knowledge_graph import KnowledgeGraph
        
        kg = KnowledgeGraph("./data/kg.db")
        result = ingest_with_entities("contract.pdf", knowledge_graph=kg)
        
        print(f"Found {len(result['extraction']['entities'])} entities")
    """
    # Ingest document
    ingestion_result = ingest_document(pdf_path, **ingest_kwargs)
    
    # Extract entities
    extraction = extract_entities_from_tree(
        tree=ingestion_result.tree,
        extract_relationships=extract_relationships,
    )
    
    result = {
        "ingestion_result": ingestion_result,
        "extraction": extraction,
        "links": [],
    }
    
    # Store in knowledge graph if provided
    if knowledge_graph is not None:
        from rnsr.extraction import EntityLinker
        
        # Store entities
        for entity in extraction["entities"]:
            knowledge_graph.add_entity(entity)
        
        # Store relationships
        for relationship in extraction["relationships"]:
            knowledge_graph.add_relationship(relationship)
        
        # Link entities if enabled
        if link_entities:
            linker = EntityLinker(knowledge_graph)
            links = linker.link_all_entities_in_document(ingestion_result.tree.id)
            result["links"] = links
        
        logger.info(
            "stored_in_knowledge_graph",
            doc_id=ingestion_result.tree.id,
            entities=len(extraction["entities"]),
            relationships=len(extraction["relationships"]),
            links=len(result["links"]),
        )
    
    return result


# =============================================================================
# Table Detection
# =============================================================================


def detect_tables_in_tree(tree: DocumentTree) -> list[DetectedTable]:
    """
    Detect tables in all nodes of a document tree.
    
    Uses TableParser to find markdown/ASCII tables in node content.
    
    Args:
        tree: The ingested DocumentTree.
        
    Returns:
        List of DetectedTable objects with table data.
    """
    from rnsr.ingestion.table_parser import TableParser
    
    parser = TableParser(infer_types=True, detect_headers=True)
    detected_tables: list[DetectedTable] = []
    
    def process_node(node: DocumentNode) -> None:
        """Process a single node for tables."""
        if node.content:
            try:
                parsed_tables = parser.parse_from_text(
                    text=node.content,
                    doc_id=tree.id,
                    page_num=node.page_num,
                    node_id=node.id,
                )
                
                for parsed in parsed_tables:
                    # Convert ParsedTable to DetectedTable (lightweight storage format)
                    data = []
                    for row in parsed.rows:
                        row_data = [cell.value for cell in row.cells]
                        data.append(row_data)
                    
                    detected = DetectedTable(
                        id=parsed.id,
                        node_id=node.id,
                        page_num=parsed.page_num,
                        title=parsed.title,
                        headers=parsed.headers,
                        num_rows=parsed.num_rows,
                        num_cols=parsed.num_cols,
                        data=data,
                    )
                    detected_tables.append(detected)
                    
            except Exception as e:
                logger.warning(
                    "table_detection_failed",
                    node_id=node.id,
                    error=str(e)[:100],
                )
        
        # Process children
        for child in node.children:
            process_node(child)
    
    process_node(tree.root)
    
    if detected_tables:
        logger.info(
            "tables_detected",
            doc_id=tree.id,
            table_count=len(detected_tables),
        )
    
    return detected_tables


def _add_tables_to_result(result: IngestionResult) -> IngestionResult:
    """
    Detect tables in the ingestion result and add them.
    
    This is called after successful ingestion to enrich the result
    with detected table data.
    """
    try:
        tables = detect_tables_in_tree(result.tree)
        result.tables = tables
        result.stats["tables_detected"] = len(tables)
    except Exception as e:
        logger.warning("table_detection_failed", error=str(e))
        result.warnings.append(f"Table detection failed: {e}")
    
    return result


# Convenience exports
__all__ = [
    "ingest_document",
    "ingest_document_enhanced",
    "ingest_with_entities",
    "extract_entities_from_tree",
    "detect_tables_in_tree",
    "IngestionResult",
]
